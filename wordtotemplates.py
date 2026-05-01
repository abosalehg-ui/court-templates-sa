# -*- coding: utf-8 -*-
"""
محول قوالب الجلسات
أداة تحويل ملفات Word (.docx) إلى صيغة templates.js

الميزات:
- سحب وإفلات + زر استعراض
- معاينة JSON قبل الحفظ مع إمكانية التعديل
- وضع ليلي/نهاري بزر يدوي مع حفظ التفضيل
- دعم الشاشات عالية الدقة
- واجهة عربية كاملة (RTL)
"""
import os
import sys
import json

# دعم الشاشات عالية الدقة (يجب وضعه قبل استيراد QtWidgets)
os.environ.setdefault("QT_AUTO_SCREEN_SCALE_FACTOR", "1")
os.environ.setdefault("QT_ENABLE_HIGHDPI_SCALING", "1")

from PyQt5.QtCore import (
    Qt, QSettings, QMimeData, QTimer, QSize,
)
from PyQt5.QtGui import (
    QFont, QFontDatabase, QIcon, QPalette, QColor, QDragEnterEvent, QDropEvent,
    QPixmap, QPainter,
)
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QFileDialog, QListWidget, QListWidgetItem,
    QTextEdit, QSplitter, QMessageBox, QFrame, QStackedWidget,
    QToolButton, QSizePolicy, QStyle,
)

from docx import Document


# ════════════════════════ منطق التحويل ════════════════════════
# (كان في ملف converter.py سابقاً، تم دمجه هنا ليصبح كل شيء في ملف واحد)

def _get_cell_text(cell):
    """يستخرج نص الخلية مع الحفاظ على فواصل الأسطر بين الفقرات."""
    paragraphs = [p.text for p in cell.paragraphs]
    return '\n'.join(paragraphs).strip('\n').strip()


def _extract_keyword_from_content(content):
    """
    عند فراغ خلية الـ keyword، نأخذ أول كلمتين عربيتين من المحتوى.
    التوقف عند أول كلمة لا تحتوي حروفاً عربية (مثل النقاط أو الأرقام).
    """
    if not content:
        return ""
    tokens = content.split()
    keyword_parts = []
    for tok in tokens[:5]:
        has_arabic = any('\u0600' <= c <= '\u06FF' for c in tok)
        if has_arabic:
            keyword_parts.append(tok)
            if len(keyword_parts) >= 2:
                break
        else:
            break
    return ' '.join(keyword_parts)


def parse_docx(filepath):
    """
    يحلل ملف docx ويعيد قائمة من العناصر:
    [{"num": "1", "keyword": "...", "content": "..."}, ...]

    المنطق:
    - يكشف تلقائياً عدد الأعمدة (2 = keyword|content، 3 = num|keyword|content)
    - يتجاهل الترويسات المدموجة والصفوف الفارغة
    - يعيد الترقيم تسلسلياً من 1
    """
    doc = Document(filepath)
    items = []
    counter = 1

    for table in doc.tables:
        for row in table.rows:
            cells = list(row.cells)
            if len(cells) < 2:
                continue

            # نأخذ "الخلايا المنطقية" بإزالة التكرارات الناتجة عن دمج الخلايا
            # في python-docx: الخلايا المدموجة تشير لنفس عنصر <w:tc>
            logical_texts = []
            seen_tcs = set()
            for c in cells:
                tc_id = id(c._tc)
                if tc_id not in seen_tcs:
                    seen_tcs.add(tc_id)
                    logical_texts.append(_get_cell_text(c))

            # ترويسة مدموجة بالكامل (خلية واحدة منطقية فقط) = نتجاهل
            if len(logical_texts) < 2:
                continue

            # كشف التخطيط:
            # عمودان => keyword | content
            # ثلاثة أعمدة => num (يُتجاهل) | keyword | content
            if len(logical_texts) == 2:
                keyword = logical_texts[0]
                content = logical_texts[1]
            else:
                keyword = logical_texts[-2]
                content = logical_texts[-1]

            # تخطي الصفوف بلا محتوى
            if not content:
                continue

            # استخراج keyword من المحتوى عند فراغه
            if not keyword:
                keyword = _extract_keyword_from_content(content)

            items.append({
                "num": str(counter),
                "keyword": keyword,
                "content": content
            })
            counter += 1

    return items


def filename_to_key(filepath):
    """'الجلسة الأولى.docx' -> 'الجلسة_الأولى'"""
    name = os.path.splitext(os.path.basename(filepath))[0]
    return name.replace(' ', '_')


def build_js_content(data):
    """يبني نص ملف JS بصيغة const templatesData = {...};"""
    json_str = json.dumps(data, ensure_ascii=False, indent=4)
    return f"const templatesData = {json_str};\n"


def parse_files(filepaths):
    """يحلل مجموعة ملفات ويعيد قاموساً: {"key1": [...], "key2": [...]}"""
    data = {}
    for fp in filepaths:
        key = filename_to_key(fp)
        data[key] = parse_docx(fp)
    return data


# ====================== الأنماط (Themes) ======================

LIGHT_THEME = {
    "bg":            "#F7F8FA",
    "bg_alt":        "#FFFFFF",
    "bg_drop":       "#EEF2FF",
    "bg_drop_hover": "#E0E7FF",
    "border":        "#D1D5DB",
    "border_focus":  "#6366F1",
    "text":          "#111827",
    "text_muted":    "#6B7280",
    "primary":       "#4F46E5",
    "primary_hover": "#4338CA",
    "primary_text":  "#FFFFFF",
    "secondary":     "#E5E7EB",
    "secondary_hover":"#D1D5DB",
    "secondary_text":"#111827",
    "danger":        "#DC2626",
    "danger_hover":  "#B91C1C",
    "success":       "#10B981",
    "success_hover": "#059669",
    "code_bg":       "#F3F4F6",
    "code_text":     "#1F2937",
    "shadow":        "rgba(0,0,0,0.08)",
}

DARK_THEME = {
    "bg":            "#0F172A",
    "bg_alt":        "#1E293B",
    "bg_drop":       "#1E293B",
    "bg_drop_hover": "#334155",
    "border":        "#334155",
    "border_focus":  "#818CF8",
    "text":          "#F1F5F9",
    "text_muted":    "#94A3B8",
    "primary":       "#6366F1",
    "primary_hover": "#818CF8",
    "primary_text":  "#FFFFFF",
    "secondary":     "#334155",
    "secondary_hover":"#475569",
    "secondary_text":"#F1F5F9",
    "danger":        "#EF4444",
    "danger_hover":  "#DC2626",
    "success":       "#10B981",
    "success_hover": "#34D399",
    "code_bg":       "#0F172A",
    "code_text":     "#E2E8F0",
    "shadow":        "rgba(0,0,0,0.4)",
}


def build_stylesheet(theme):
    """يبني ورقة الأنماط بناءً على القاموس المعطى."""
    return f"""
    QMainWindow, QWidget#mainContainer {{
        background-color: {theme['bg']};
        color: {theme['text']};
    }}

    QWidget {{
        color: {theme['text']};
        font-family: 'Segoe UI', 'Tajawal', 'Cairo', Arial, sans-serif;
    }}

    /* ====== شريط العنوان ====== */
    QFrame#headerBar {{
        background-color: {theme['bg_alt']};
        border: none;
        border-bottom: 1px solid {theme['border']};
    }}

    QLabel#appTitle {{
        font-size: 20px;
        font-weight: 700;
        color: {theme['text']};
        padding: 4px 8px;
    }}

    QLabel#appSubtitle {{
        font-size: 12px;
        color: {theme['text_muted']};
        padding: 0 8px;
    }}

    /* ====== زر تبديل الوضع ====== */
    QToolButton#themeToggle {{
        background-color: {theme['secondary']};
        color: {theme['secondary_text']};
        border: 1px solid {theme['border']};
        border-radius: 8px;
        padding: 8px 14px;
        font-size: 13px;
        font-weight: 600;
        min-width: 100px;
    }}
    QToolButton#themeToggle:hover {{
        background-color: {theme['secondary_hover']};
    }}

    /* ====== منطقة السحب والإفلات ====== */
    QFrame#dropZone {{
        background-color: {theme['bg_drop']};
        border: 2px dashed {theme['border']};
        border-radius: 12px;
        padding: 20px;
    }}
    QFrame#dropZone[active="true"] {{
        background-color: {theme['bg_drop_hover']};
        border: 2px dashed {theme['border_focus']};
    }}
    QLabel#dropLabel {{
        font-size: 16px;
        font-weight: 600;
        color: {theme['text']};
        background: transparent;
    }}
    QLabel#dropHint {{
        font-size: 13px;
        color: {theme['text_muted']};
        background: transparent;
    }}

    /* ====== الأزرار ====== */
    QPushButton {{
        background-color: {theme['secondary']};
        color: {theme['secondary_text']};
        border: 1px solid {theme['border']};
        border-radius: 8px;
        padding: 9px 18px;
        font-size: 13px;
        font-weight: 600;
        min-height: 20px;
    }}
    QPushButton:hover {{
        background-color: {theme['secondary_hover']};
    }}
    QPushButton:disabled {{
        opacity: 0.5;
        color: {theme['text_muted']};
    }}

    QPushButton#primaryBtn {{
        background-color: {theme['primary']};
        color: {theme['primary_text']};
        border: 1px solid {theme['primary']};
    }}
    QPushButton#primaryBtn:hover {{
        background-color: {theme['primary_hover']};
        border: 1px solid {theme['primary_hover']};
    }}
    QPushButton#primaryBtn:disabled {{
        background-color: {theme['secondary']};
        color: {theme['text_muted']};
        border: 1px solid {theme['border']};
    }}

    QPushButton#successBtn {{
        background-color: {theme['success']};
        color: white;
        border: 1px solid {theme['success']};
    }}
    QPushButton#successBtn:hover {{
        background-color: {theme['success_hover']};
        border: 1px solid {theme['success_hover']};
    }}

    QPushButton#dangerBtn {{
        background-color: transparent;
        color: {theme['danger']};
        border: 1px solid {theme['danger']};
    }}
    QPushButton#dangerBtn:hover {{
        background-color: {theme['danger']};
        color: white;
    }}

    /* ====== قائمة الملفات ====== */
    QListWidget {{
        background-color: {theme['bg_alt']};
        color: {theme['text']};
        border: 1px solid {theme['border']};
        border-radius: 8px;
        padding: 6px;
        font-size: 13px;
        outline: none;
    }}
    QListWidget::item {{
        background-color: transparent;
        padding: 8px 12px;
        border-radius: 6px;
        margin: 2px 0;
    }}
    QListWidget::item:hover {{
        background-color: {theme['secondary']};
    }}
    QListWidget::item:selected {{
        background-color: {theme['primary']};
        color: {theme['primary_text']};
    }}

    /* ====== محرر النص (المعاينة) ====== */
    QTextEdit {{
        background-color: {theme['code_bg']};
        color: {theme['code_text']};
        border: 1px solid {theme['border']};
        border-radius: 8px;
        padding: 12px;
        font-family: 'Consolas', 'Cascadia Code', 'Courier New', monospace;
        font-size: 13px;
        selection-background-color: {theme['primary']};
        selection-color: {theme['primary_text']};
    }}
    QTextEdit:focus {{
        border: 1px solid {theme['border_focus']};
    }}

    /* ====== التسميات ====== */
    QLabel#sectionLabel {{
        font-size: 14px;
        font-weight: 700;
        color: {theme['text']};
        padding: 4px 0;
    }}
    QLabel#statusLabel {{
        font-size: 12px;
        color: {theme['text_muted']};
        padding: 4px 8px;
    }}
    QLabel#statusLabel[type="success"] {{
        color: {theme['success']};
        font-weight: 600;
    }}
    QLabel#statusLabel[type="error"] {{
        color: {theme['danger']};
        font-weight: 600;
    }}

    /* ====== الفاصل (Splitter) ====== */
    QSplitter::handle {{
        background-color: {theme['border']};
        width: 4px;
    }}
    QSplitter::handle:hover {{
        background-color: {theme['border_focus']};
    }}

    /* ====== شريط التمرير ====== */
    QScrollBar:vertical {{
        background: transparent;
        width: 12px;
        border: none;
    }}
    QScrollBar::handle:vertical {{
        background: {theme['border']};
        border-radius: 6px;
        min-height: 30px;
    }}
    QScrollBar::handle:vertical:hover {{
        background: {theme['text_muted']};
    }}
    QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {{
        background: none;
        border: none;
        height: 0;
    }}
    QScrollBar:horizontal {{
        background: transparent;
        height: 12px;
        border: none;
    }}
    QScrollBar::handle:horizontal {{
        background: {theme['border']};
        border-radius: 6px;
        min-width: 30px;
    }}
    QScrollBar::handle:horizontal:hover {{
        background: {theme['text_muted']};
    }}
    QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {{
        background: none;
        border: none;
        width: 0;
    }}

    /* ====== صناديق الحوار ====== */
    QMessageBox {{
        background-color: {theme['bg_alt']};
        color: {theme['text']};
    }}
    QMessageBox QLabel {{
        color: {theme['text']};
        font-size: 13px;
    }}
    QMessageBox QPushButton {{
        min-width: 80px;
        padding: 6px 16px;
    }}
    """


# ====================== الواجهة الرئيسية ======================

class DropZone(QFrame):
    """منطقة السحب والإفلات."""
    def __init__(self, on_files_dropped, parent=None):
        super().__init__(parent)
        self.on_files_dropped = on_files_dropped
        self.setObjectName("dropZone")
        self.setProperty("active", "false")
        self.setAcceptDrops(True)
        self.setMinimumHeight(120)

        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignCenter)
        layout.setSpacing(8)

        self.label = QLabel("اسحب ملفات Word هنا")
        self.label.setObjectName("dropLabel")
        self.label.setAlignment(Qt.AlignCenter)

        self.hint = QLabel("أو انقر على زر «استعراض» أدناه — يدعم ملفات .docx متعددة")
        self.hint.setObjectName("dropHint")
        self.hint.setAlignment(Qt.AlignCenter)

        layout.addWidget(self.label)
        layout.addWidget(self.hint)

    def _set_active(self, active):
        self.setProperty("active", "true" if active else "false")
        # إعادة تطبيق النمط لتفعيل الـ property selector
        self.style().unpolish(self)
        self.style().polish(self)

    def dragEnterEvent(self, event: QDragEnterEvent):
        mime: QMimeData = event.mimeData()
        if mime.hasUrls():
            for url in mime.urls():
                if url.toLocalFile().lower().endswith('.docx'):
                    event.acceptProposedAction()
                    self._set_active(True)
                    return
        event.ignore()

    def dragLeaveEvent(self, event):
        self._set_active(False)
        super().dragLeaveEvent(event)

    def dropEvent(self, event: QDropEvent):
        self._set_active(False)
        urls = event.mimeData().urls()
        files = []
        for url in urls:
            path = url.toLocalFile()
            if path.lower().endswith('.docx'):
                files.append(path)
        if files:
            self.on_files_dropped(files)
            event.acceptProposedAction()


class MainWindow(QMainWindow):
    SETTINGS_ORG = "DocxToJsTool"
    SETTINGS_APP = "TemplatesConverter"

    def __init__(self):
        super().__init__()
        self.settings = QSettings(self.SETTINGS_ORG, self.SETTINGS_APP)
        self.is_dark = self.settings.value("darkMode", False, type=bool)

        # حالة التطبيق
        self.files = []          # قائمة بمسارات الملفات المختارة
        self.parsed_data = None  # نتيجة التحليل (قاموس)

        self._setup_ui()
        self._apply_theme()

    # ---------- بناء الواجهة ----------
    def _setup_ui(self):
        self.setWindowTitle("محول قوالب الجلسات — Word إلى JavaScript")
        self.setMinimumSize(1100, 720)
        self.setLayoutDirection(Qt.RightToLeft)

        central = QWidget()
        central.setObjectName("mainContainer")
        self.setCentralWidget(central)

        root = QVBoxLayout(central)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # شريط العنوان
        root.addWidget(self._build_header())

        # المحتوى الرئيسي
        body = QWidget()
        body_layout = QVBoxLayout(body)
        body_layout.setContentsMargins(20, 16, 20, 20)
        body_layout.setSpacing(14)

        # Splitter أفقي يفصل بين قسم الملفات وقسم المعاينة
        splitter = QSplitter(Qt.Horizontal)
        # في RTL، نضع المعاينة على اليسار والملفات على اليمين بشكل طبيعي
        splitter.addWidget(self._build_files_panel())
        splitter.addWidget(self._build_preview_panel())
        splitter.setStretchFactor(0, 0)
        splitter.setStretchFactor(1, 1)
        splitter.setSizes([380, 700])

        body_layout.addWidget(splitter, 1)

        # شريط الحالة السفلي
        body_layout.addWidget(self._build_status_bar())

        root.addWidget(body, 1)

    def _build_header(self):
        header = QFrame()
        header.setObjectName("headerBar")
        header.setFixedHeight(64)

        h = QHBoxLayout(header)
        h.setContentsMargins(20, 8, 20, 8)
        h.setSpacing(12)

        # شعار + عنوان
        title_box = QVBoxLayout()
        title_box.setSpacing(0)

        title = QLabel("⚖️  محول قوالب الجلسات")
        title.setObjectName("appTitle")
        subtitle = QLabel("Word → templates.js")
        subtitle.setObjectName("appSubtitle")

        title_box.addWidget(title)
        title_box.addWidget(subtitle)

        h.addLayout(title_box)
        h.addStretch(1)

        # زر تبديل الوضع
        self.theme_btn = QToolButton()
        self.theme_btn.setObjectName("themeToggle")
        self.theme_btn.clicked.connect(self._toggle_theme)
        self._update_theme_btn()

        h.addWidget(self.theme_btn)
        return header

    def _build_files_panel(self):
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(10)

        # عنوان القسم
        title = QLabel("📂  الملفات")
        title.setObjectName("sectionLabel")
        layout.addWidget(title)

        # منطقة السحب والإفلات
        self.drop_zone = DropZone(self._add_files)
        layout.addWidget(self.drop_zone)

        # أزرار الإضافة والمسح
        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)

        self.browse_btn = QPushButton("📁  استعراض الملفات")
        self.browse_btn.clicked.connect(self._browse_files)
        btn_row.addWidget(self.browse_btn)

        self.clear_btn = QPushButton("🗑️  مسح الكل")
        self.clear_btn.setObjectName("dangerBtn")
        self.clear_btn.clicked.connect(self._clear_files)
        btn_row.addWidget(self.clear_btn)

        layout.addLayout(btn_row)

        # قائمة الملفات
        self.files_list = QListWidget()
        self.files_list.setSelectionMode(QListWidget.ExtendedSelection)
        self.files_list.setMinimumHeight(180)
        layout.addWidget(self.files_list, 1)

        # زر إزالة المحدد
        self.remove_btn = QPushButton("➖  إزالة المحدد")
        self.remove_btn.clicked.connect(self._remove_selected)
        layout.addWidget(self.remove_btn)

        # زر التحويل
        self.convert_btn = QPushButton("🔄  تحويل ومعاينة")
        self.convert_btn.setObjectName("primaryBtn")
        self.convert_btn.clicked.connect(self._convert)
        self.convert_btn.setMinimumHeight(40)
        layout.addWidget(self.convert_btn)

        return panel

    def _build_preview_panel(self):
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(10)

        # شريط أدوات المعاينة
        tools = QHBoxLayout()
        tools.setSpacing(8)

        title = QLabel("📄  معاينة الإخراج (templates.js)")
        title.setObjectName("sectionLabel")
        tools.addWidget(title)
        tools.addStretch(1)

        self.edit_btn = QPushButton("✏️  تفعيل التعديل")
        self.edit_btn.setCheckable(True)
        self.edit_btn.clicked.connect(self._toggle_edit)
        self.edit_btn.setEnabled(False)
        tools.addWidget(self.edit_btn)

        self.copy_btn = QPushButton("📋  نسخ")
        self.copy_btn.clicked.connect(self._copy_to_clipboard)
        self.copy_btn.setEnabled(False)
        tools.addWidget(self.copy_btn)

        self.save_btn = QPushButton("💾  حفظ كـ .js")
        self.save_btn.setObjectName("successBtn")
        self.save_btn.clicked.connect(self._save_file)
        self.save_btn.setEnabled(False)
        self.save_btn.setMinimumHeight(36)
        tools.addWidget(self.save_btn)

        layout.addLayout(tools)

        # محرر المعاينة
        self.preview = QTextEdit()
        self.preview.setReadOnly(True)
        self.preview.setLineWrapMode(QTextEdit.NoWrap)
        self.preview.setLayoutDirection(Qt.LeftToRight)
        # خط مونوسبيس للكود
        mono_font = QFont("Consolas", 11)
        mono_font.setStyleHint(QFont.Monospace)
        self.preview.setFont(mono_font)
        self.preview.setPlaceholderText(
            "اختر ملفات Word ثم اضغط «تحويل ومعاينة»\nستظهر النتيجة هنا قبل الحفظ، ويمكن تعديلها يدوياً"
        )

        layout.addWidget(self.preview, 1)
        return panel

    def _build_status_bar(self):
        self.status_label = QLabel("جاهز — اسحب ملفات .docx أو استعرضها للبدء")
        self.status_label.setObjectName("statusLabel")
        self.status_label.setProperty("type", "default")
        return self.status_label

    # ---------- إجراءات الملفات ----------
    def _browse_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "اختر ملفات Word",
            "",
            "ملفات Word (*.docx);;كل الملفات (*)"
        )
        if files:
            self._add_files(files)

    def _add_files(self, paths):
        added = 0
        for p in paths:
            if not p.lower().endswith('.docx'):
                continue
            if p in self.files:
                continue
            self.files.append(p)
            added += 1
            item = QListWidgetItem(f"📄  {os.path.basename(p)}")
            item.setData(Qt.UserRole, p)
            item.setToolTip(p)
            self.files_list.addItem(item)
        if added:
            self._set_status(f"تمت إضافة {added} ملف — الإجمالي: {len(self.files)}")

    def _remove_selected(self):
        items = self.files_list.selectedItems()
        if not items:
            self._set_status("لم يتم تحديد أي ملف للإزالة", level="error")
            return
        for it in items:
            path = it.data(Qt.UserRole)
            if path in self.files:
                self.files.remove(path)
            self.files_list.takeItem(self.files_list.row(it))
        self._set_status(f"تمت الإزالة — الإجمالي: {len(self.files)}")

    def _clear_files(self):
        if not self.files:
            return
        ok = QMessageBox.question(
            self,
            "تأكيد المسح",
            f"هل تريد مسح {len(self.files)} ملف من القائمة؟",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        if ok == QMessageBox.Yes:
            self.files.clear()
            self.files_list.clear()
            self._set_status("تم مسح القائمة")

    # ---------- التحويل والمعاينة ----------
    def _convert(self):
        if not self.files:
            QMessageBox.warning(
                self,
                "لا توجد ملفات",
                "الرجاء إضافة ملف واحد على الأقل قبل التحويل."
            )
            return

        self._set_status("جاري التحويل…")
        QApplication.processEvents()

        try:
            self.parsed_data = parse_files(self.files)
            js_text = build_js_content(self.parsed_data)
            self.preview.setPlainText(js_text)
            # تفعيل الأزرار
            self.save_btn.setEnabled(True)
            self.copy_btn.setEnabled(True)
            self.edit_btn.setEnabled(True)

            total_items = sum(len(v) for v in self.parsed_data.values())
            sessions = len(self.parsed_data)
            self._set_status(
                f"✓ تم التحويل بنجاح — {sessions} جلسة، {total_items} عنصر إجمالاً",
                level="success"
            )
        except Exception as e:
            QMessageBox.critical(
                self,
                "خطأ في التحويل",
                f"حدث خطأ أثناء قراءة الملفات:\n\n{str(e)}"
            )
            self._set_status(f"خطأ: {str(e)}", level="error")

    def _toggle_edit(self):
        is_editing = self.edit_btn.isChecked()
        self.preview.setReadOnly(not is_editing)
        if is_editing:
            self.edit_btn.setText("🔒  قفل التعديل")
            self._set_status("وضع التعديل: مفعّل — يمكنك تعديل المعاينة قبل الحفظ")
        else:
            self.edit_btn.setText("✏️  تفعيل التعديل")
            self._set_status("وضع التعديل: مقفل")

    def _copy_to_clipboard(self):
        text = self.preview.toPlainText()
        if not text:
            return
        QApplication.clipboard().setText(text)
        self._set_status("✓ تم نسخ المحتوى إلى الحافظة", level="success")

    def _save_file(self):
        text = self.preview.toPlainText()
        if not text:
            QMessageBox.warning(self, "لا يوجد محتوى", "لا توجد معاينة قابلة للحفظ.")
            return

        path, _ = QFileDialog.getSaveFileName(
            self,
            "حفظ الملف",
            "templates.js",
            "ملفات JavaScript (*.js);;كل الملفات (*)"
        )
        if not path:
            return

        try:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(text)
            self._set_status(f"✓ تم الحفظ في: {path}", level="success")
            QMessageBox.information(
                self,
                "تم الحفظ",
                f"تم حفظ الملف بنجاح في:\n\n{path}"
            )
        except Exception as e:
            QMessageBox.critical(
                self,
                "خطأ في الحفظ",
                f"تعذّر حفظ الملف:\n\n{str(e)}"
            )

    # ---------- الوضع الليلي/النهاري ----------
    def _toggle_theme(self):
        self.is_dark = not self.is_dark
        self.settings.setValue("darkMode", self.is_dark)
        self._apply_theme()
        self._update_theme_btn()

    def _apply_theme(self):
        theme = DARK_THEME if self.is_dark else LIGHT_THEME
        self.setStyleSheet(build_stylesheet(theme))

    def _update_theme_btn(self):
        if self.is_dark:
            self.theme_btn.setText("☀️  الوضع النهاري")
        else:
            self.theme_btn.setText("🌙  الوضع الليلي")

    # ---------- شريط الحالة ----------
    def _set_status(self, msg, level="default"):
        self.status_label.setText(msg)
        self.status_label.setProperty("type", level)
        self.status_label.style().unpolish(self.status_label)
        self.status_label.style().polish(self.status_label)


# ====================== التشغيل ======================

def setup_high_dpi():
    """يهيّئ دعم الشاشات عالية الدقة قبل إنشاء QApplication."""
    if hasattr(Qt, 'AA_EnableHighDpiScaling'):
        QApplication.setAttribute(Qt.AA_EnableHighDpiScaling, True)
    if hasattr(Qt, 'AA_UseHighDpiPixmaps'):
        QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps, True)


def main():
    setup_high_dpi()
    app = QApplication(sys.argv)
    app.setApplicationName("محول قوالب الجلسات")
    app.setOrganizationName("DocxToJsTool")

    # خط افتراضي مناسب للعربية
    default_font = QFont("Segoe UI", 10)
    app.setFont(default_font)

    win = MainWindow()
    win.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
